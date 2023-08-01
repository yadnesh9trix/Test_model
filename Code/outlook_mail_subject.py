import datetime
import win32com.client as win32
import pandas as pd
from docx import Document

def get_outlook_inbox_subjects():
    # Connect to Outlook Application
    outlook_app = win32.Dispatch("Outlook.Application").GetNamespace("MAPI")

    # Get Inbox folder
    inbox = outlook_app.GetDefaultFolder(6)  # 6 represents the Inbox folder

    # Get all the emails in the Inbox
    # emails = inbox.Items

    # Filter emails received today
    start_date = "2023-07-31"  # Replace with your desired start date and time (00:00:00 for the start of the day)
    end_date = "2023-07-31"

    # Use the Items collection and filter emails within the specified date range
    start_datetime = datetime.datetime.strptime(start_date, '%Y-%m-%d')
    end_datetime = datetime.datetime.strptime(end_date, '%Y-%m-%d')
    # Format dates as mm/dd/yyyy hh:mm AM/PM
    # Format datetime objects to match ReceivedTime format
    start_datetime_str = start_datetime.strftime('%d/%m/%Y %H:%M %p')
    end_datetime_str = end_datetime.strftime('%d/%m/%Y 11:59 PM')

    # Filter emails using ReceivedTime with timezone information
    emails = inbox.Items.Restrict("[ReceivedTime] >= '" + start_datetime_str + "' AND [ReceivedTime] <= '" + end_datetime_str + "'")

    # Create a list to store email subjects
    # data = []
    # for item in emails:
    #     subject = item.Subject
    #     date_received = item.ReceivedTime.strftime('%Y-%m-%d %H:%M:%S')
    #     content = item.Body  # Get the content of the email
    #     data.append({'Subject': subject, 'Date Received': date_received, 'Content': content})
    return emails


def save_emails_to_word(emails, file_path):
    doc = Document()

    for item in emails:
        subject = item.Subject
        date_received = item.ReceivedTime.strftime('%Y-%m-%d %H:%M:%S')
        body = item.Body.strip()

        # Add subject and date to the Word document
        # body = body.replace('\n', '')
        body = body.replace('\r', ' ')
        doc.add_heading(subject, level=0)
        doc.add_paragraph("Received Date: " + date_received)
        doc.add_paragraph(body)

        # Add a page break between emails
        doc.add_page_break()

    # Save the Word document
    doc.save(file_path)

def save_subjects_to_excel(data, file_path):
    # Create a DataFrame with email subjects
    df = pd.DataFrame(data ,columns=['Subject','Date Received'])

    # Save DataFrame to Excel file
    df.to_excel(file_path, index=False)


if __name__ == "__main__":
    # Replace 'C:/path/to/output/emails.xlsx' with your desired file path
    output_file_path = 'D:/emails.docx'

    # Fetch email subjects
    outlook_emails = get_outlook_inbox_subjects()

    # Save subjects to Excel
    # save_subjects_to_excel(email_subjects, output_file_path)

    # Save emails to Word document
    save_emails_to_word(outlook_emails, output_file_path)